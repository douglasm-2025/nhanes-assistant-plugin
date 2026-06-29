# server.py — NHANES Assistant MCP server (path-derived; works unchanged from nhanes_server/ or nhanes_plugin/)
from fastmcp import FastMCP
import subprocess
import os
import re
import json
import zipfile
import tempfile
from datetime import datetime

mcp = FastMCP("nhanes-assistant")

SCRIPT_DIR           = os.path.dirname(os.path.abspath(__file__))
PLUGIN_ROOT          = os.path.dirname(SCRIPT_DIR)
R_HELPERS_DIR        = os.path.join(SCRIPT_DIR, "r_helpers")
NHANES_STYLES_PATH   = os.path.join(R_HELPERS_DIR, "nhanes_styles.R")
NHANES_SURVEY_PATH   = os.path.join(R_HELPERS_DIR, "nhanes_survey.R")
STROBE_NUT_PATH      = os.path.join(PLUGIN_ROOT, "STROBE-nut_checklist.docx")
TEMPLATE_PATH        = os.path.join(SCRIPT_DIR, "templates", "manuscript_template.docx")
ANALYSIS_SCRIPTS_DIR = os.path.join(SCRIPT_DIR, "analysis_scripts")
HTML_DIR             = os.path.join(SCRIPT_DIR, "outputs", "html")
DOCX_DIR             = os.path.join(SCRIPT_DIR, "outputs", "docx")
FIGURES_DIR          = os.path.join(SCRIPT_DIR, "outputs", "figures")
CSV_DIR              = os.path.join(SCRIPT_DIR, "outputs", "csv")
BUNDLES_DIR          = os.path.join(SCRIPT_DIR, "outputs", "bundles")

for _d in [ANALYSIS_SCRIPTS_DIR, HTML_DIR, DOCX_DIR, FIGURES_DIR, CSV_DIR, BUNDLES_DIR]:
    os.makedirs(_d, exist_ok=True)


def _load_strobe_nut() -> str:
    if not os.path.exists(STROBE_NUT_PATH):
        return (
            f"STROBE-nut checklist not found at {STROBE_NUT_PATH}. "
            "Ensure STROBE-nut_checklist.docx is in the plugin root and run setup.sh."
        )
    try:
        import docx as _docx
        doc = _docx.Document(STROBE_NUT_PATH)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"Error loading STROBE-nut checklist: {e}"


_STROBE_NUT_CONTENT = _load_strobe_nut()

# Applied after theme_nhanes() when a figure is rendered with blank_axes=True
# (diagram-style figures such as participant flow charts have no meaningful axes).
_BLANK_AXES_R = (
    "p <- p + ggplot2::theme(axis.line = ggplot2::element_blank(), "
    "axis.text = ggplot2::element_blank(), axis.ticks = ggplot2::element_blank(), "
    "axis.title = ggplot2::element_blank())"
)

_HTML_STYLE_CSS = """
body { font-family: Arial, sans-serif; font-size: 11pt; max-width: 920px; margin: 0 auto; padding: 24px; color: #333; line-height: 1.5; }
h2 { font-size: 13pt; font-weight: bold; margin-top: 28px; }
.nhanes-figure { margin: 24px 0; }
.nhanes-figure .figure-title { font-weight: bold; font-size: 11pt; margin-bottom: 6px; }
.nhanes-figure figcaption.figure-caption { font-style: italic; font-size: 9pt; color: #555; margin-top: 6px; }
.nhanes-table { margin: 24px 0; }
.nhanes-table .table-title { font-weight: bold; font-size: 11pt; margin-bottom: 6px; }
.nhanes-table .table-notes { font-style: italic; font-size: 9pt; color: #555; margin-top: 6px; }
table { border-collapse: collapse; width: 100%; font-size: 11pt; }
th { background-color: #333333; color: white; font-weight: bold; padding: 8px 12px; text-align: left; }
td { padding: 6px 12px; border-bottom: 1px solid #e8e8e8; }
tr:nth-child(even) td { background-color: #f2f2f2; }
"""

# ─────────────────────────────────────────────────────────────
# Tools — Core
# ─────────────────────────────────────────────────────────────

@mcp.tool()
def search_nhanes_variables(query: str) -> str:
    """Call nhanesA::nhanesSearch() to find NHANES variable codes matching a text query."""
    r_code = r"""
library(nhanesA)
query <- Sys.getenv("NHANES_QUERY")
tryCatch({
  res <- nhanesSearch(query, ignore.case = TRUE)
  if (is.data.frame(res) && nrow(res) > 0) {
    cols <- intersect(c("Variable.Name","Variable.Description","Data.File.Name"), names(res))
    print(res[, cols])
  } else {
    cat("No results found for query:", query, "\n")
  }
}, error = function(e) cat("Search error:", conditionMessage(e), "\n"))
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.R', delete=False) as f:
        f.write(r_code)
        tmp = f.name
    try:
        env = {**os.environ, 'NHANES_QUERY': query}
        result = subprocess.run(
            ["Rscript", "--vanilla", tmp],
            capture_output=True, text=True, env=env, timeout=600
        )
        if result.returncode == 0:
            return result.stdout
        return f"Search failed.\nStderr:\n{result.stderr}"
    except subprocess.TimeoutExpired:
        return "Error: nhanesSearch timed out after 600 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


@mcp.tool()
def execute_r_script(r_code: str) -> str:
    """Save r_code to a timestamped file in analysis_scripts/ and execute it via Rscript.

    The canonical survey helpers (r_helpers/nhanes_survey.R: nh_design, nh_mean,
    nh_diff) are auto-sourced before the script runs, so analyses can call them
    without an explicit source() and produce correct, near-identical survey
    boilerplate every run.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(ANALYSIS_SCRIPTS_DIR, f"analysis_{ts}.R")
    survey_fwd = NHANES_SURVEY_PATH.replace("\\", "/")
    preamble = (
        "# --- NHANES Assistant: canonical survey helpers (auto-loaded) ---\n"
        f'if (file.exists("{survey_fwd}")) source("{survey_fwd}")\n'
        "# --- end auto-load ---\n\n"
    )
    with open(path, 'w') as f:
        f.write(preamble + r_code)
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", path],
            capture_output=True, text=True, timeout=600
        )
        if result.returncode == 0:
            return f"Script saved to: {path}\n\n{result.stdout}"
        return (
            f"Script saved to: {path}\n\n"
            f"Execution failed.\nStderr:\n{result.stderr}\n\nStdout:\n{result.stdout}"
        )
    except subprocess.TimeoutExpired:
        return f"Script saved to: {path}\n\nError: timed out after 600 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."


@mcp.tool()
def write_csv_outputs(csv_string: str, filename: str) -> str:
    """Write csv_string to outputs/csv/<filename>."""
    if not filename.endswith(".csv"):
        filename += ".csv"
    out = os.path.join(CSV_DIR, os.path.basename(filename))
    try:
        with open(out, 'w') as f:
            f.write(csv_string)
        return f"CSV written to: {out}"
    except Exception as e:
        return f"Failed to write CSV: {e}"


@mcp.tool()
def lookup_nhanes_codebook(table_name: str, cycle: str) -> str:
    """Return all variable names and descriptions for a NHANES table+cycle via nhanesA::nhanesTableVars()."""
    # Validate inputs before interpolating into R source (injection guard).
    if not re.fullmatch(r"[A-Za-z0-9_]+", table_name) or not re.fullmatch(r"[A-Za-z0-9_]+", cycle):
        return "Invalid table_name or cycle: only letters, digits, and underscores are allowed."
    tbl_full = f"{table_name}_{cycle}"
    r_code = f"""
library(nhanesA)
tryCatch({{
  vars <- nhanesTableVars("{tbl_full}")
  print(vars)
}}, error = function(e) cat("Error:", conditionMessage(e), "\\n"))
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.R', delete=False) as f:
        f.write(r_code)
        tmp = f.name
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", tmp],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0:
            return result.stdout
        return f"Lookup failed.\nStderr:\n{result.stderr}"
    except subprocess.TimeoutExpired:
        return "Error: codebook lookup timed out after 120 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


# ─────────────────────────────────────────────────────────────
# Tools — HTML rendering
# ─────────────────────────────────────────────────────────────

@mcp.tool()
def render_html_figure(r_code: str, title: str, caption: str, blank_axes: bool = False) -> str:
    """
    Execute r_code producing ggplot2 object `p`. Applies theme_nhanes() automatically,
    renders to SVG at 7×5 in, wraps in <figure> HTML. Saves to outputs/html/. Returns HTML fragment.
    Set blank_axes=True for diagram-style figures (e.g. participant flow charts) that should have
    no axis lines, ticks, text, or titles — it is applied AFTER theme_nhanes() so it is not overridden.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    svg_path    = os.path.join(FIGURES_DIR, f"fig_{ts}.svg")
    script_path = os.path.join(ANALYSIS_SCRIPTS_DIR, f"fig_{ts}.R")
    styles_fwd  = NHANES_STYLES_PATH.replace("\\", "/")
    svg_fwd     = svg_path.replace("\\", "/")

    wrapped = f"""library(ggplot2)
library(svglite)
source("{styles_fwd}")

{r_code}

p <- p + theme_nhanes()
{_BLANK_AXES_R if blank_axes else ""}
svglite::svglite("{svg_fwd}", width = 7, height = 5)
print(p)
dev.off()
cat("SVG written\\n")
"""
    with open(script_path, 'w') as f:
        f.write(wrapped)
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", script_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            return f"Figure rendering failed.\nStderr:\n{result.stderr}\n\nStdout:\n{result.stdout}"
        with open(svg_path, 'r') as f:
            svg_content = f.read()
        fragment = (
            f'<figure class="nhanes-figure">\n'
            f'  <p class="figure-title"><strong>{title}</strong></p>\n'
            f'  {svg_content}\n'
            f'  <figcaption class="figure-caption"><em>{caption}</em></figcaption>\n'
            f'</figure>'
        )
        with open(os.path.join(HTML_DIR, f"fig_{ts}.html"), 'w') as f:
            f.write(fragment)
        return fragment
    except subprocess.TimeoutExpired:
        return "Error: figure rendering timed out after 120 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."


@mcp.tool()
def render_html_table(r_code: str, title: str, notes: str) -> str:
    """
    Execute r_code producing gt table object `tbl`. Applies nhanes_table_style() automatically.
    Returns HTML fragment. Saves to outputs/html/.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    script_path = os.path.join(ANALYSIS_SCRIPTS_DIR, f"tbl_{ts}.R")
    styles_fwd  = NHANES_STYLES_PATH.replace("\\", "/")

    wrapped = f"""library(gt)
source("{styles_fwd}")

{r_code}

tbl <- nhanes_table_style(tbl)
html_str <- gt::as_raw_html(tbl)
cat(html_str)
"""
    with open(script_path, 'w') as f:
        f.write(wrapped)
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", script_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            return f"Table rendering failed.\nStderr:\n{result.stderr}"
        fragment = (
            f'<div class="nhanes-table">\n'
            f'  <p class="table-title"><strong>{title}</strong></p>\n'
            f'  {result.stdout}\n'
            f'  <p class="table-notes"><em>{notes}</em></p>\n'
            f'</div>'
        )
        with open(os.path.join(HTML_DIR, f"tbl_{ts}.html"), 'w') as f:
            f.write(fragment)
        return fragment
    except subprocess.TimeoutExpired:
        return "Error: table rendering timed out after 120 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."


@mcp.tool()
def write_html_output(html_content: str, filename: str) -> str:
    """Wrap html_content in a standalone HTML document with embedded CSS. Saves to outputs/html/."""
    if not filename.endswith(".html"):
        filename += ".html"
    out = os.path.join(HTML_DIR, os.path.basename(filename))
    full = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>NHANES Analysis Output</title>
  <style>
{_HTML_STYLE_CSS}
  </style>
</head>
<body>
{html_content}
</body>
</html>"""
    with open(out, 'w') as f:
        f.write(full)
    return out


# ─────────────────────────────────────────────────────────────
# Tools — Manuscript rendering (PNG + RDS)
# ─────────────────────────────────────────────────────────────

@mcp.tool()
def render_manuscript_figure(r_code: str, title: str, caption: str, figure_number: int,
                             blank_axes: bool = False) -> str:
    """
    Execute r_code producing ggplot2 object `p`. Applies theme_nhanes(). Renders to PNG at 300 DPI
    via ragg::agg_png(). Saves to outputs/figures/figN_TIMESTAMP.png. Returns file path.
    Set blank_axes=True for diagram-style figures (e.g. participant flow charts) that should have
    no axis lines, ticks, text, or titles — it is applied AFTER theme_nhanes() so it is not overridden.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    png_path    = os.path.join(FIGURES_DIR, f"fig{figure_number}_{ts}.png")
    script_path = os.path.join(ANALYSIS_SCRIPTS_DIR, f"fig{figure_number}_{ts}.R")
    styles_fwd  = NHANES_STYLES_PATH.replace("\\", "/")
    png_fwd     = png_path.replace("\\", "/")

    wrapped = f"""library(ggplot2)
library(ragg)
source("{styles_fwd}")

{r_code}

p <- p + theme_nhanes()
{_BLANK_AXES_R if blank_axes else ""}
ragg::agg_png("{png_fwd}", width = 7, height = 5, units = "in", res = 300)
print(p)
dev.off()
cat("PNG written\\n")
"""
    with open(script_path, 'w') as f:
        f.write(wrapped)
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", script_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            return f"Manuscript figure rendering failed.\nStderr:\n{result.stderr}"
        return png_path
    except subprocess.TimeoutExpired:
        return "Error: manuscript figure rendering timed out after 120 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."


@mcp.tool()
def render_manuscript_table(r_code: str, title: str, notes: str, table_number: int) -> str:
    """
    Execute r_code producing flextable object `ft`. Applies nhanes_flextable_style().
    Serializes to RDS at outputs/figures/tblN_TIMESTAMP.rds. CSV backup in outputs/csv/.
    Returns RDS file path for write_docx_output.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    rds_path    = os.path.join(FIGURES_DIR, f"tbl{table_number}_{ts}.rds")
    csv_path    = os.path.join(CSV_DIR, f"tbl{table_number}_{ts}.csv")
    script_path = os.path.join(ANALYSIS_SCRIPTS_DIR, f"tbl{table_number}_{ts}.R")
    styles_fwd  = NHANES_STYLES_PATH.replace("\\", "/")
    rds_fwd     = rds_path.replace("\\", "/")
    csv_fwd     = csv_path.replace("\\", "/")

    wrapped = f"""library(flextable)
library(officer)
source("{styles_fwd}")

{r_code}

ft <- nhanes_flextable_style(ft)
saveRDS(ft, "{rds_fwd}")
tryCatch({{
  body_data <- ft$body$dataset
  if (!is.null(body_data)) write.csv(body_data, "{csv_fwd}", row.names = FALSE)
}}, error = function(e) invisible(NULL))
cat("Table RDS saved to: {rds_fwd}\\n")
"""
    with open(script_path, 'w') as f:
        f.write(wrapped)
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", script_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            return f"Manuscript table rendering failed.\nStderr:\n{result.stderr}"
        return rds_path
    except subprocess.TimeoutExpired:
        return "Error: manuscript table rendering timed out after 120 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."


# ─────────────────────────────────────────────────────────────
# Tools — Docx assembly and bundle
# ─────────────────────────────────────────────────────────────

@mcp.tool()
def write_docx_output(
    title: str,
    abstract: str,
    methods: str,
    results: str,
    limitations: str,
    ethics_statement: str,
    data_availability: str,
    figure_paths: list,
    table_rds_paths: list,
    authors: str = "",
    figure_legends: list = None,
    table_legends: list = None
) -> str:
    """
    Assemble a complete .docx manuscript. Reads manuscript_template.docx via officer.
    Embeds 300 DPI PNGs from figure_paths and flextables from table_rds_paths.
    authors is optional — user fills in the delivered .docx manually.
    figure_legends/table_legends are optional lists aligned by position with
    figure_paths/table_rds_paths: when supplied, each figure gets a legend below
    it ("Figure N. <legend>") and each table a title above it ("Table N. <legend>").
    Returns output file path.
    """
    if figure_legends is None:
        figure_legends = []
    if table_legends is None:
        table_legends = []
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(DOCX_DIR, f"manuscript_{ts}.docx")

    content = {
        "title": title, "authors": authors, "abstract": abstract,
        "methods": methods, "results": results, "limitations": limitations,
        "ethics_statement": ethics_statement, "data_availability": data_availability,
        "figure_paths": figure_paths, "table_rds_paths": table_rds_paths,
        "figure_legends": figure_legends, "table_legends": table_legends,
        "output_path": output_path.replace("\\", "/"),
        "template_path": TEMPLATE_PATH.replace("\\", "/"),
    }

    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as jf:
        json.dump(content, jf, ensure_ascii=False)
        json_path = jf.name

    r_code = f"""
library(officer)
library(flextable)
library(jsonlite)

content <- fromJSON("{json_path.replace(chr(92), '/')}", simplifyVector = TRUE)

doc <- read_docx(content$template_path)

add_paras <- function(doc, text, style = "Normal") {{
  if (is.null(text) || !nchar(trimws(text))) return(doc)
  for (p in Filter(nchar, trimws(strsplit(text, "\\n\\n")[[1]])))
    doc <- body_add_par(doc, p, style = style)
  doc
}}

doc <- body_add_par(doc, content$title, style = "Normal")
if (nchar(content$authors)) doc <- body_add_par(doc, content$authors, style = "Normal")
doc <- body_add_par(doc, "Abstract",         style = "heading 1")
doc <- add_paras(doc, content$abstract)
doc <- body_add_par(doc, "Methods",          style = "heading 1")
doc <- add_paras(doc, content$methods)
doc <- body_add_par(doc, "Results",          style = "heading 1")
doc <- add_paras(doc, content$results)
doc <- body_add_par(doc, "Limitations",      style = "heading 1")
doc <- add_paras(doc, content$limitations)
doc <- body_add_par(doc, "Ethics Statement", style = "heading 1")
doc <- add_paras(doc, content$ethics_statement)
doc <- body_add_par(doc, "Data Availability",style = "heading 1")
doc <- add_paras(doc, content$data_availability)

leg_at <- function(vec, i) {{
  if (length(vec) >= i && !is.na(vec[[i]]) && nchar(trimws(vec[[i]]))) trimws(vec[[i]]) else ""
}}

if (length(content$figure_paths) > 0) {{
  doc <- body_add_par(doc, "Figure legends", style = "heading 1")
  for (i in seq_along(content$figure_paths)) {{
    fp <- content$figure_paths[[i]]
    if (file.exists(fp)) {{
      doc <- body_add_img(doc, src = fp, width = 6, height = 4.3)
      leg <- leg_at(content$figure_legends, i)
      cap <- if (nchar(leg)) sprintf("Figure %d. %s", i, leg) else sprintf("Figure %d.", i)
      doc <- body_add_par(doc, cap, style = "Normal")  # legend below the figure
    }}
  }}
}}

if (length(content$table_rds_paths) > 0) {{
  doc <- body_add_par(doc, "Tables", style = "heading 1")
  for (i in seq_along(content$table_rds_paths)) {{
    rds_p <- content$table_rds_paths[[i]]
    if (file.exists(rds_p)) {{
      leg <- leg_at(content$table_legends, i)
      cap <- if (nchar(leg)) sprintf("Table %d. %s", i, leg) else sprintf("Table %d.", i)
      doc <- body_add_par(doc, cap, style = "Normal")  # title above the table
      ft <- readRDS(rds_p)
      doc <- body_add_flextable(doc, ft)
      doc <- body_add_par(doc, "", style = "Normal")
    }}
  }}
}}

print(doc, target = content$output_path)
cat("DOCX saved to:", content$output_path, "\\n")
"""

    with tempfile.NamedTemporaryFile(mode='w', suffix='.R', delete=False) as rf:
        rf.write(r_code)
        r_path = rf.name
    try:
        result = subprocess.run(
            ["Rscript", "--vanilla", r_path],
            capture_output=True, text=True, timeout=300
        )
        if result.returncode != 0:
            return f"DOCX assembly failed.\nStderr:\n{result.stderr}\n\nStdout:\n{result.stdout}"
        return output_path
    except subprocess.TimeoutExpired:
        return "Error: DOCX assembly timed out after 300 seconds."
    except OSError as e:
        return f"Error: could not run Rscript ({e}). Ensure R is installed and on PATH (run setup.sh)."
    finally:
        for p in [json_path, r_path]:
            if os.path.exists(p):
                os.remove(p)


@mcp.tool()
def create_reproducibility_bundle(session_timestamp: str = "") -> str:
    """
    Zip the most recent .docx, all R scripts, all CSVs, and all HTML files into outputs/bundles/.
    Returns the zip file path.

    session_timestamp: optional timestamp prefix (e.g. "20240615_143022") produced by a single
    analysis session. When provided and non-empty, only analysis_scripts/, csv/, and html/ files
    whose filename *contains* that prefix are included — stale artifacts from prior runs are
    excluded. The most-recent .docx selection is unaffected. When empty, all files are included
    (backward-compatible behavior).
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = os.path.join(BUNDLES_DIR, f"nhanes_results_{ts}.zip")

    def _include(filename: str) -> bool:
        """Return True if the file should be included given session_timestamp filter."""
        if not session_timestamp:
            return True
        return session_timestamp in filename

    to_zip = []
    docx_files = sorted(
        [os.path.join(DOCX_DIR, f) for f in os.listdir(DOCX_DIR) if f.endswith(".docx")],
        key=os.path.getmtime, reverse=True
    )
    if docx_files:
        to_zip.append(("docx/" + os.path.basename(docx_files[0]), docx_files[0]))
    for f in os.listdir(ANALYSIS_SCRIPTS_DIR):
        if f.endswith(".R") and _include(f):
            to_zip.append(("scripts/" + f, os.path.join(ANALYSIS_SCRIPTS_DIR, f)))
    for f in os.listdir(CSV_DIR):
        if f.endswith(".csv") and _include(f):
            to_zip.append(("csv/" + f, os.path.join(CSV_DIR, f)))
    for f in os.listdir(HTML_DIR):
        if f.endswith(".html") and _include(f):
            to_zip.append(("html/" + f, os.path.join(HTML_DIR, f)))

    if not to_zip:
        return "No output files found to bundle. Run an analysis first."

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for arcname, filepath in to_zip:
            if os.path.exists(filepath):
                zf.write(filepath, arcname)
    return zip_path


# ─────────────────────────────────────────────────────────────
# Resources
# ─────────────────────────────────────────────────────────────

@mcp.resource("config://nhanes_expertise")
def get_nhanes_expertise() -> str:
    return """NHANES Expert Reference
=======================

SURVEY CYCLES — include all except K; flag L
A=1999-2000 | B=2001-2002 | C=2003-2004 | D=2005-2006 | E=2007-2008 | F=2009-2010
G=2011-2012 | H=2013-2014 | I=2015-2016 | J=2017-2018
K=2019-2020 — EXCLUDE ALWAYS: COVID-19 disrupted data collection (non-representative)
L=2021-2022 — FLAG: partial-year data; include only with explicit acknowledgment

FILE NAMING: TABLENAME_SUFFIX e.g. DEMO_J, BMX_J, DR1TOT_J, DR1IFF_J
Common: DEMO (demographics), BMX (body measures), BPX/BPXO (blood pressure 2017+),
DR1TOT/DR2TOT (dietary totals Day 1/2), DR1IFF/DR2IFF (individual foods),
PAQ (physical activity), ALQ (alcohol), SMQ (smoking), DIQ (diabetes),
MCQ (medical conditions), BPQ (blood pressure questionnaire)

CROSS-CYCLE HARMONIZATION — always check before pooling cycles:
1. RACE/ETHNICITY: RIDRETH1 (cycles A-F) → RIDRETH3 (cycles G+, adds NH Asian).
   Cannot pool across cycle F/G boundary without harmonization.
2. DIETARY FORMAT: Consistent DR1IFF/DR1TOT structure starts cycle B (2001-02).
   Avoid cycle A for dietary analyses.
3. COTININE: Measurement method changed in cycle H (2013-14). Flag if analyzing
   cotinine across the G/H boundary.
4. BLOOD PRESSURE: BPX (cycles A-I) → BPXO (cycles J+). Use correct table per cycle.

SURVEY DESIGN VARIABLES
PSU: SDMVPSU | Strata: SDMVSTRA | Always nest=TRUE
Set options(survey.lonely.psu="adjust") BEFORE any svydesign() call.

WEIGHT SELECTION — use weight for the component with lowest probability of selection:
Interview/questionnaire only         → WTINT2YR
MEC exam / physical measures / labs  → WTMEC2YR
Day 1 dietary recall                 → WTDRD1  (in DR1TOT and DR1IFF tables)
Day 1 + Day 2 dietary recall         → WTDR2D  — MUST filter WTDR2D > 0 before svydesign()
Fasting biochemistry subsample       → WTSAF2YR
Environmental chemicals subsample    → WTSB2YR

COMBINED-CYCLE WEIGHT RECALCULATION (required when pooling N two-year cycles):
  df$WTDRD1_combined <- df$WTDRD1 / n_cycles   # before svydesign()

KEY DEMOGRAPHICS
RIDAGEYR  — age in years (continuous)
RIAGENDR  — sex: 1=Male, 2=Female
RIDRETH3  — race/ethnicity: 1=Mexican American, 2=Other Hispanic, 3=NH White,
             4=NH Black, 6=NH Asian, 7=Other/Multi-racial
INDFMPIR  — poverty income ratio (0-5+; higher = wealthier)
DMDEDUC2  — education (adults 20+)
RIDEXPRG  — pregnant: 1=Yes, 2=No, 3=Cannot determine

BMI VARIABLES
Adult: BMXBMI (continuous kg/m², age 20+)
Pediatric: BMDBMIC (factor with TEXT labels — ages 2-19, CDC growth chart categories)
  Labels: "Underweight", "Normal weight", "Overweight", "Obese"
  CRITICAL: nhanesA returns BMDBMIC as a factor with text labels, NOT integers.
  Use as.character(BMDBMIC) %in% c("Overweight","Obese") — NEVER BMDBMIC >= 3.

DIETARY VARIABLES
DR1_030Z  — eating occasion. nhanesA returns TEXT labels (not integer codes):
  Snack labels: "Snack", "Merienda", "Botana", "Bocadillo", "Tentempie"
  Meal labels: "Breakfast","Desayano","Lunch","Almuerzo","Dinner","Comida","Cena","Brunch"
DR1ILINE  — food item line number. Use arrange(SEQN, DR1ILINE) for within-person ordering.
DR1TKCAL  — total energy (kcal, Day 1)
DR1TPROT  — protein (g) | DR1TTFAT — total fat (g) | DR1TCARB — carbohydrate (g)
DR1TSODI  — sodium (mg)

GOLDBERG MISREPORTER DETECTION (dietary analyses)
Schofield equation BMR (kcal/day) by sex+age, then EI/BMR ratio.
Implausible range: < 0.64 (under-reporter) or > 2.72 (over-reporter) — Goldberg cutoffs.
Report n excluded; offer as sensitivity analysis if n_excluded > 5% of analytic sample.

ENERGY ADJUSTMENT METHODS (when nutrient is the exposure)
Residual method: svyglm(nutrient ~ total_energy, design=svy); use Pearson residuals as adjusted nutrient.
Nutrient density: nutrient / DR1TKCAL * 1000 (per 1000 kcal).
State the chosen method in Methods and justify (nut-12.2).
"""


@mcp.resource("config://strobe_nut")
def get_strobe_nut() -> str:
    return _STROBE_NUT_CONTENT


@mcp.resource("config://reporting_conventions")
def get_reporting_conventions() -> str:
    return """Statistical Reporting Conventions
==================================
Effect estimates: 2 decimal places  (e.g., OR = 1.23, β = -0.45)
95% CIs in parentheses:  (X.XX, X.XX)  — not brackets, not ±
P-values: 3 decimal places; floor at <0.001 — never write p=0.000
Weighted means: X.XX ± SE  or  X.XX (95% CI: X.XX, X.XX)
Prevalence: X.X% (95% CI: X.X%, X.X%)
Survey design reported: PSU variable name, strata variable name, weight variable, software+version
Counts ≥1,000: format with commas (e.g., n = 3,214 not n = 3214)
Rounding: apply only to final reported values; never round intermediate calculations
Reference categories must be stated explicitly in text (e.g., "Non-Hispanic White was the reference")
Subgroup n: report both unweighted n and weighted % for Table 1 cells
"""


@mcp.resource("config://html_style")
def get_html_style() -> str:
    return _HTML_STYLE_CSS


@mcp.resource("config://nhanes_styles")
def get_nhanes_styles() -> str:
    """Return the full R source of nhanes_styles.R for agents to reference."""
    if os.path.exists(NHANES_STYLES_PATH):
        with open(NHANES_STYLES_PATH, 'r') as f:
            return f.read()
    return "nhanes_styles.R not found. Run setup.sh first."


@mcp.resource("config://nhanes_survey")
def get_nhanes_survey() -> str:
    """Return the full R source of nhanes_survey.R (canonical survey helpers) for agents to reference."""
    if os.path.exists(NHANES_SURVEY_PATH):
        with open(NHANES_SURVEY_PATH, 'r') as f:
            return f.read()
    return "nhanes_survey.R not found. Run setup.sh first."


if __name__ == "__main__":
    mcp.run()
